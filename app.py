import streamlit as st
import pandas as pd
import math
from io import BytesIO
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.utils import get_column_letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4, landscape


# ==========================================================
# BUSINESS LOGIC (Balanced Distribution)
# ==========================================================

def generate_slot_duty(rooms, teachers, slots):

    if not teachers:
        return {}, {}, 0

    total_duties = len(rooms) * slots
    base_duty = total_duties // len(teachers)
    extra = total_duties % len(teachers)

    # Teachers get either base_duty or base_duty + 1
    duty_limit = {
        teacher: base_duty + (1 if i < extra else 0)
        for i, teacher in enumerate(teachers)
    }

    duty_table = {room: [""] * slots for room in rooms}
    teacher_count = {t: 0 for t in teachers}
    slot_assignments = {slot: [] for slot in range(slots)}

    teacher_index = 0

    for room in rooms:
        for slot in range(slots):

            attempts = 0
            while attempts < len(teachers):

                teacher = teachers[teacher_index % len(teachers)]
                teacher_index += 1
                attempts += 1

                if (
                    teacher_count[teacher] < duty_limit[teacher]
                    and teacher not in slot_assignments[slot]
                    and teacher not in duty_table[room]
                ):
                    duty_table[room][slot] = teacher
                    teacher_count[teacher] += 1
                    slot_assignments[slot].append(teacher)
                    break

            if duty_table[room][slot] == "":
                duty_table[room][slot] = "No Available Teacher"

    return duty_table, teacher_count, max(duty_limit.values())


# ==========================================================
# EXPORT FUNCTIONS
# ==========================================================

def export_excel(df):
    buffer = BytesIO()

    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Duty List")
        worksheet = writer.sheets["Duty List"]

        for i, column in enumerate(df.columns, 1):
            max_length = max(
                df[column].astype(str).map(len).max(),
                len(column)
            )
            worksheet.column_dimensions[get_column_letter(i)].width = max_length + 3

    buffer.seek(0)
    return buffer


def export_word(df, header_lines):
    buffer = BytesIO()
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    for line in header_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(" ")

    table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
    table.style = "Table Grid"

    for col in range(df.shape[1]):
        table.rows[0].cells[col].text = str(df.columns[col])

    for row in range(df.shape[0]):
        for col in range(df.shape[1]):
            table.rows[row + 1].cells[col].text = str(df.iat[row, col])

    doc.add_paragraph("\nPRINCIPAL")
    doc.add_paragraph(f"Generated on: {pd.Timestamp.now().date()}")

    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_pdf(df, header_lines):
    buffer = BytesIO()

    pdf = SimpleDocTemplate(buffer, pagesize=landscape(A4))
    elements = []
    style = getSampleStyleSheet()

    for line in header_lines:
        elements.append(Paragraph(f"<b>{line}</b>", style["Title"]))
        elements.append(Spacer(1, 6))

    elements.append(Spacer(1, 12))

    data = [df.columns.tolist()] + df.values.tolist()
    col_width = (landscape(A4)[0] - 40) / len(df.columns)

    table = Table(data, colWidths=[col_width] * len(df.columns), repeatRows=1)

    table.setStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
    ])

    elements.append(table)
    elements.append(Spacer(1, 20))
    elements.append(Paragraph("PRINCIPAL", style["Normal"]))

    pdf.build(elements)
    buffer.seek(0)
    return buffer


# ==========================================================
# STREAMLIT UI
# ==========================================================

st.set_page_config(layout="wide")
st.title("🤖 AI-Based Examination Duty Allocation System")

school_name = st.text_input("School Name")
school_address = st.text_input("School Address")
exam_title = st.text_input("Exam Title")
duty_date = st.date_input("Duty Date")

num_slots = st.number_input("Number of Slots", min_value=1, max_value=10)

slot_timings = []
for i in range(num_slots):
    start = st.text_input(f"Slot {i+1} Start", key=f"s{i}")
    end = st.text_input(f"Slot {i+1} End", key=f"e{i}")
    slot_timings.append(f"{start}-{end}" if start and end else "")

teachers_input = st.text_area("Teacher Names (comma separated)")
rooms_input = st.text_area("Room Names (comma separated)")

if st.button("Generate Duty Plan"):

    teachers = [t.strip() for t in teachers_input.split(",") if t.strip()]
    rooms = [r.strip() for r in rooms_input.split(",") if r.strip()]

    result, teacher_count, max_duty = generate_slot_duty(rooms, teachers, num_slots)

    df = pd.DataFrame(result).T
    df.columns = [f"Slot {i+1} ({slot_timings[i]})" for i in range(num_slots)]
    df.insert(0, "Room", df.index)
    df.insert(0, "S.No", range(1, len(df) + 1))
    df.reset_index(drop=True, inplace=True)

    st.success(f"Max Duties per Teacher: {max_duty}")
    st.dataframe(df, width="stretch")

    # Fairness Check
    counts = list(teacher_count.values())
    if counts and max(counts) - min(counts) <= 1:
        st.success("✔ Duties distributed almost equally.")
    else:
        st.warning("⚠ Duties not evenly distributed.")

    header = [
        school_name,
        school_address,
        f"Duty List for {duty_date} – {exam_title}"
    ]

    st.download_button("📥 Excel", export_excel(df), "Duty_List.xlsx")
    st.download_button("📄 Word", export_word(df, header), "Duty_List.docx")
    st.download_button("📑 PDF", export_pdf(df, header), "Duty_List.pdf")